"""Scan local folders for evidence and map it to controls.

Two scans:
  scan_documents(folder) -> list of Chunk   (pdf, docx, xlsx, md, txt, csv, json, yaml)
  scan_environment(folder) -> list of Signal (config/code artefacts that hint at a control)
Then retrieve(chunks, control) -> top chunks for a control, BM25-ranked.
Everything runs locally; nothing is sent anywhere until assess() is called.
"""
from __future__ import annotations

import fnmatch
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path

DOC_EXT = {".pdf", ".docx", ".xlsx", ".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".log"}
SKIP_DIRS = {".git", "node_modules", ".venv", "venv", "__pycache__", ".idea", ".vscode", "site-packages"}
MAX_FILE_MB = 25
CHUNK = 1800
OVERLAP = 200


@dataclass
class Chunk:
    path: str
    text: str
    idx: int

    @property
    def label(self):
        return f"{Path(self.path).name} [part {self.idx + 1}]"


@dataclass
class Signal:
    path: str
    kind: str
    hint: str
    controls: list


# ---------- text extraction ----------
def _pdf(p: Path) -> str:
    from pypdf import PdfReader
    r = PdfReader(str(p))
    return "\n".join((pg.extract_text() or "") for pg in r.pages)


def _docx(p: Path) -> str:
    import docx
    d = docx.Document(str(p))
    parts = [para.text for para in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx(p: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(p), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"## Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i > 2000:
                break
            vals = [str(v) for v in row if v is not None]
            if vals:
                out.append(" | ".join(vals))
    return "\n".join(out)


def extract(p: Path) -> str:
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _pdf(p)
        if ext == ".docx":
            return _docx(p)
        if ext == ".xlsx":
            return _xlsx(p)
        return p.read_text(errors="ignore")
    except Exception as e:  # unreadable file is a finding, not a crash
        return f"[unreadable: {e}]"


def chunk_text(text: str) -> list[str]:
    text = re.sub(r"[ \t]+", " ", text)
    out, i = [], 0
    while i < len(text):
        out.append(text[i:i + CHUNK])
        i += CHUNK - OVERLAP
    return [c for c in out if c.strip()]


def iter_files(folder: str, exts: set[str]):
    for p in Path(folder).rglob("*"):
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        if p.is_file() and p.suffix.lower() in exts and p.stat().st_size < MAX_FILE_MB * 1e6:
            yield p


def scan_documents(folder: str, progress=None) -> list[Chunk]:
    chunks = []
    files = list(iter_files(folder, DOC_EXT))
    for n, p in enumerate(files):
        if progress:
            progress(n, len(files), p.name)
        for i, c in enumerate(chunk_text(f"[{p.name}]\n" + extract(p))):
            chunks.append(Chunk(str(p), c, i))
    return chunks


# ---------- environment signals ----------
# (glob, kind, hint, control ids that this artefact is relevant to)
ENV_RULES = [
    ("*.tf", "iac", "Infrastructure-as-code present — deployment is codified and reviewable", ["M3", "S3.1", "A.6"]),
    ("Dockerfile", "container", "Container build definition — check for pinned base images and non-root user", ["M3", "S3.1"]),
    ("docker-compose*.yml", "container", "Container orchestration config", ["M3"]),
    (".github/workflows/*.yml", "ci", "CI pipeline — check for tests, scans and approval gates before deploy", ["M3", "D3.2", "S3.2"]),
    ("*.gitlab-ci.yml", "ci", "CI pipeline — check for tests, scans and approval gates before deploy", ["M3", "D3.2", "S3.2"]),
    ("*model_card*", "model-doc", "Model card — documentation of intended use, limits and evaluation", ["D1.1", "M2", "A.6.2", "MAP 1.1"]),
    ("*MODEL_CARD*", "model-doc", "Model card — documentation of intended use, limits and evaluation", ["D1.1", "M2", "A.6.2"]),
    ("*eval*", "evaluation", "Evaluation artefacts — tests or golden sets for model behaviour", ["D3.1", "M4", "MEASURE 2"]),
    ("*test*", "evaluation", "Test artefacts", ["D3.1", "M4"]),
    ("*logging*", "logging", "Logging configuration — check what agent actions are captured and retained", ["S4.1", "D4.1", "M5"]),
    ("*log4j*", "logging", "Logging configuration", ["S4.1", "D4.1"]),
    ("*iam*", "access", "IAM/access policy — check least privilege for agent identities", ["S1.2", "S2.1", "D2.1"]),
    ("*rbac*", "access", "Role-based access config", ["S1.2", "D2.1"]),
    ("*.env", "secrets", "Environment file — secrets may be stored in plaintext; must not be in evidence or repos", ["S1.3", "D2.2"]),
    ("*secret*", "secrets", "Secrets-related file — verify it is a reference, not plaintext credentials", ["S1.3", "D2.2"]),
    ("*prompt*", "prompt", "Prompt/system-prompt files — governed as configuration? versioned? reviewed?", ["D2.3", "S2.2", "M2"]),
    ("*guardrail*", "guardrail", "Guardrail configuration — policy-bound execution evidence", ["S2.1", "D2.3"]),
    ("*policy*.json", "policy", "Machine-readable policy — check who can change it and how changes are reviewed", ["S2.1", "S1.1"]),
    ("*incident*", "incident", "Incident records or runbooks", ["D4.3", "S4.3", "M6"]),
    ("*runbook*", "incident", "Runbooks — operational response procedures", ["D4.3", "S4.3"]),
    ("*inventory*", "inventory", "Inventory register — agent/model registration evidence", ["S1.1", "M1.3", "A.9"]),
    ("*register*", "inventory", "Register document", ["S1.1", "M1.3"]),
    ("*dpia*", "privacy", "Data protection impact assessment", ["M2", "A.8"]),
    ("*pdpa*", "privacy", "PDPA-related document", ["M2", "A.8"]),
    ("requirements*.txt", "dependencies", "Python dependency manifest — third-party AI libraries and pinned versions", ["M3", "A.10"]),
    ("package.json", "dependencies", "Node dependency manifest", ["M3", "A.10"]),
    ("*sbom*", "dependencies", "Software bill of materials", ["M3", "A.10"]),
]


def scan_environment(folder: str) -> list[Signal]:
    out, seen = [], set()
    for p in Path(folder).rglob("*"):
        if any(part in SKIP_DIRS for part in p.parts) or not p.is_file():
            continue
        rel = str(p.relative_to(folder))
        for pat, kind, hint, ctls in ENV_RULES:
            if fnmatch.fnmatch(p.name.lower(), pat.lower()) or fnmatch.fnmatch(rel.lower(), pat.lower()):
                if (rel, kind) not in seen:
                    seen.add((rel, kind))
                    out.append(Signal(rel, kind, hint, ctls))
                break
    return out


# ---------- retrieval ----------
_tok = lambda s: re.findall(r"[a-z0-9][a-z0-9\-\.]{1,}", s.lower())


class Index:
    def __init__(self, chunks: list[Chunk]):
        from rank_bm25 import BM25Okapi
        self.chunks = chunks
        self.bm25 = BM25Okapi([_tok(c.text) for c in chunks]) if chunks else None

    def query(self, control, k=3, min_score=3.0):
        if not self.bm25:
            return []
        q = _tok(f"{control.title} {control.req} {control.maps}")
        scores = self.bm25.get_scores(q)
        top = sorted(range(len(scores)), key=lambda i: -scores[i])[:k]
        return [(self.chunks[i], float(scores[i])) for i in top if scores[i] >= min_score]


def signals_for(control, signals: list[Signal]) -> list[Signal]:
    cid = control.id.split()[0].rstrip("★").strip()
    fam = cid.split(".")[0]
    return [s for s in signals if any(c == cid or c == fam or cid.startswith(c + ".") for c in s.controls)]
